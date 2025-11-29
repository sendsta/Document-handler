#!/usr/bin/env python3
"""
Tender Documents Handler MCP Server
====================================
A Model Context Protocol (MCP) server for parsing, analyzing, importing,
OCR processing, and managing tender documents for tri-tender.

This MCP ensures that tender documents uploaded by users are accessible,
readable, and usable for generating better tender responses.
"""

import asyncio
import base64
import hashlib
import json
import logging
import mimetypes
import os
import re
import shutil
import subprocess
import tempfile
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Any, Optional

# MCP SDK imports
from mcp.server import Server
from mcp.server.stdio import stdio_server
from mcp.types import (
    CallToolResult,
    ListToolsResult,
    TextContent,
    Tool,
)

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("tender-docs-mcp")

# ============================================================================
# Configuration
# ============================================================================

UPLOAD_DIR = Path(os.environ.get("TENDER_UPLOAD_DIR", "/tmp/tender-uploads"))
PROCESSED_DIR = Path(os.environ.get("TENDER_PROCESSED_DIR", "/tmp/tender-processed"))
CACHE_DIR = Path(os.environ.get("TENDER_CACHE_DIR", "/tmp/tender-cache"))

# Ensure directories exist
for dir_path in [UPLOAD_DIR, PROCESSED_DIR, CACHE_DIR]:
    dir_path.mkdir(parents=True, exist_ok=True)


class DocumentType(Enum):
    PDF = "pdf"
    DOCX = "docx"
    DOC = "doc"
    TXT = "txt"
    HTML = "html"
    RTF = "rtf"
    ODT = "odt"
    IMAGE = "image"
    UNKNOWN = "unknown"


# ============================================================================
# Data Classes
# ============================================================================

@dataclass
class DocumentMetadata:
    """Metadata extracted from a document."""
    file_name: str
    file_path: str
    file_size: int
    file_type: str
    page_count: Optional[int] = None
    title: Optional[str] = None
    author: Optional[str] = None
    created_date: Optional[str] = None
    modified_date: Optional[str] = None
    word_count: Optional[int] = None
    checksum: Optional[str] = None

    def to_dict(self) -> dict:
        return {k: v for k, v in self.__dict__.items() if v is not None}


@dataclass
class ExtractedSection:
    """A section extracted from a document."""
    title: str
    content: str
    page_number: Optional[int] = None
    section_number: Optional[str] = None
    level: int = 1

    def to_dict(self) -> dict:
        return self.__dict__


@dataclass
class TenderRequirement:
    """A requirement extracted from a tender document."""
    requirement_id: str
    description: str
    category: str
    is_mandatory: bool = False
    page_number: Optional[int] = None
    source_text: Optional[str] = None

    def to_dict(self) -> dict:
        return self.__dict__


@dataclass
class AnalysisResult:
    """Result of document analysis."""
    document_id: str
    metadata: DocumentMetadata
    sections: list[ExtractedSection] = field(default_factory=list)
    requirements: list[TenderRequirement] = field(default_factory=list)
    tables: list[dict] = field(default_factory=list)
    deadlines: list[dict] = field(default_factory=list)
    key_contacts: list[dict] = field(default_factory=list)
    evaluation_criteria: list[dict] = field(default_factory=list)
    compliance_items: list[dict] = field(default_factory=list)
    full_text: str = ""
    summary: str = ""

    def to_dict(self) -> dict:
        return {
            "document_id": self.document_id,
            "metadata": self.metadata.to_dict(),
            "sections": [s.to_dict() for s in self.sections],
            "requirements": [r.to_dict() for r in self.requirements],
            "tables": self.tables,
            "deadlines": self.deadlines,
            "key_contacts": self.key_contacts,
            "evaluation_criteria": self.evaluation_criteria,
            "compliance_items": self.compliance_items,
            "full_text": self.full_text[:5000] + "..." if len(self.full_text) > 5000 else self.full_text,
            "summary": self.summary,
        }


# ============================================================================
# Document Processing Functions
# ============================================================================

def detect_document_type(file_path: Path) -> DocumentType:
    """Detect the type of a document based on extension and MIME type."""
    extension = file_path.suffix.lower()
    
    type_mapping = {
        ".pdf": DocumentType.PDF,
        ".docx": DocumentType.DOCX,
        ".doc": DocumentType.DOC,
        ".txt": DocumentType.TXT,
        ".html": DocumentType.HTML,
        ".htm": DocumentType.HTML,
        ".rtf": DocumentType.RTF,
        ".odt": DocumentType.ODT,
        ".png": DocumentType.IMAGE,
        ".jpg": DocumentType.IMAGE,
        ".jpeg": DocumentType.IMAGE,
        ".tiff": DocumentType.IMAGE,
        ".tif": DocumentType.IMAGE,
        ".bmp": DocumentType.IMAGE,
    }
    
    return type_mapping.get(extension, DocumentType.UNKNOWN)


def calculate_checksum(file_path: Path) -> str:
    """Calculate MD5 checksum of a file."""
    hash_md5 = hashlib.md5()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_md5.update(chunk)
    return hash_md5.hexdigest()


def extract_text_from_pdf(file_path: Path) -> tuple[str, int]:
    """Extract text from a PDF file using pdfplumber."""
    try:
        import pdfplumber
        
        text_parts = []
        page_count = 0
        
        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        return "\n\n".join(text_parts), page_count
    except ImportError:
        # Fallback to pypdf
        from pypdf import PdfReader
        
        reader = PdfReader(file_path)
        text_parts = []
        for page in reader.pages:
            text_parts.append(page.extract_text())
        
        return "\n\n".join(text_parts), len(reader.pages)


def extract_text_from_docx(file_path: Path) -> str:
    """Extract text from a DOCX file using pandoc."""
    try:
        result = subprocess.run(
            ["pandoc", str(file_path), "-t", "plain"],
            capture_output=True,
            text=True,
            check=True
        )
        return result.stdout
    except subprocess.CalledProcessError:
        # Fallback to python-docx
        try:
            from docx import Document
            doc = Document(file_path)
            return "\n\n".join([para.text for para in doc.paragraphs])
        except ImportError:
            return ""


def extract_text_from_html(file_path: Path) -> str:
    """Extract text from an HTML file."""
    try:
        from bs4 import BeautifulSoup
        
        with open(file_path, "r", encoding="utf-8") as f:
            soup = BeautifulSoup(f.read(), "html.parser")
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            return soup.get_text(separator="\n")
    except ImportError:
        # Basic fallback
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
            # Simple tag removal
            content = re.sub(r'<[^>]+>', '', content)
            return content


def perform_ocr(file_path: Path) -> str:
    """Perform OCR on an image or scanned PDF."""
    try:
        import pytesseract
        from PIL import Image
        
        doc_type = detect_document_type(file_path)
        
        if doc_type == DocumentType.IMAGE:
            image = Image.open(file_path)
            return pytesseract.image_to_string(image)
        elif doc_type == DocumentType.PDF:
            from pdf2image import convert_from_path
            
            images = convert_from_path(file_path)
            text_parts = []
            for i, image in enumerate(images):
                text_parts.append(f"--- Page {i+1} ---\n")
                text_parts.append(pytesseract.image_to_string(image))
            return "\n".join(text_parts)
        else:
            return ""
    except ImportError as e:
        logger.warning(f"OCR dependencies not available: {e}")
        return ""


def extract_tables_from_pdf(file_path: Path) -> list[dict]:
    """Extract tables from a PDF document."""
    tables = []
    try:
        import pdfplumber
        
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                page_tables = page.extract_tables()
                for table_idx, table in enumerate(page_tables):
                    if table and len(table) > 1:
                        # First row as headers
                        headers = table[0] if table[0] else [f"Col_{i}" for i in range(len(table[1]))]
                        rows = table[1:]
                        tables.append({
                            "page": page_num,
                            "table_index": table_idx,
                            "headers": headers,
                            "rows": rows,
                            "row_count": len(rows)
                        })
    except Exception as e:
        logger.warning(f"Error extracting tables: {e}")
    
    return tables


def extract_metadata(file_path: Path) -> DocumentMetadata:
    """Extract metadata from a document."""
    file_stat = file_path.stat()
    doc_type = detect_document_type(file_path)
    
    metadata = DocumentMetadata(
        file_name=file_path.name,
        file_path=str(file_path),
        file_size=file_stat.st_size,
        file_type=doc_type.value,
        checksum=calculate_checksum(file_path),
        modified_date=datetime.fromtimestamp(file_stat.st_mtime).isoformat()
    )
    
    if doc_type == DocumentType.PDF:
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            metadata.page_count = len(reader.pages)
            if reader.metadata:
                metadata.title = reader.metadata.title
                metadata.author = reader.metadata.author
                if reader.metadata.creation_date:
                    metadata.created_date = str(reader.metadata.creation_date)
        except Exception as e:
            logger.warning(f"Error reading PDF metadata: {e}")
    
    elif doc_type == DocumentType.DOCX:
        try:
            from docx import Document
            doc = Document(file_path)
            core_props = doc.core_properties
            metadata.title = core_props.title
            metadata.author = core_props.author
            if core_props.created:
                metadata.created_date = core_props.created.isoformat()
        except Exception as e:
            logger.warning(f"Error reading DOCX metadata: {e}")
    
    return metadata


# ============================================================================
# Analysis Functions
# ============================================================================

def extract_sections(text: str) -> list[ExtractedSection]:
    """Extract sections from document text based on common patterns."""
    sections = []
    
    # Common section patterns
    patterns = [
        # Numbered sections (1., 1.1, etc.)
        r'^(\d+(?:\.\d+)*\.?\s*)([A-Z][^.\n]{2,}[^\n]*)',
        # Letter sections (A., B., etc.)
        r'^([A-Z]\.\s*)([A-Z][^.\n]{2,}[^\n]*)',
        # Article/Section headers
        r'^((?:ARTICLE|SECTION|PART)\s+[\dIVXivx]+[:\.\s]*)(.*?)$',
        # All caps headers
        r'^([A-Z][A-Z\s]{4,})$',
    ]
    
    lines = text.split('\n')
    current_section = None
    current_content = []
    
    for line in lines:
        is_header = False
        for pattern in patterns:
            match = re.match(pattern, line.strip(), re.MULTILINE)
            if match:
                # Save previous section
                if current_section:
                    current_section.content = '\n'.join(current_content).strip()
                    if current_section.content:
                        sections.append(current_section)
                
                # Start new section
                section_num = match.group(1).strip() if len(match.groups()) > 1 else ""
                title = match.group(2).strip() if len(match.groups()) > 1 else match.group(1).strip()
                
                current_section = ExtractedSection(
                    title=title,
                    content="",
                    section_number=section_num,
                    level=1 if '.' not in section_num else section_num.count('.') + 1
                )
                current_content = []
                is_header = True
                break
        
        if not is_header and current_section:
            current_content.append(line)
    
    # Don't forget the last section
    if current_section:
        current_section.content = '\n'.join(current_content).strip()
        if current_section.content:
            sections.append(current_section)
    
    return sections


def extract_requirements(text: str) -> list[TenderRequirement]:
    """Extract requirements from tender document text."""
    requirements = []
    req_id = 1
    
    # Patterns that indicate requirements
    requirement_patterns = [
        (r'(?:must|shall|should|required to|is required|mandatory)\s+(.{20,200})', True),
        (r'(?:the contractor|the bidder|the vendor|the supplier)\s+(?:must|shall|should)\s+(.{20,200})', True),
        (r'(?:minimum requirement|essential requirement|mandatory requirement)[:\s]+(.{20,200})', True),
        (r'(?:criterion|criteria)[:\s]+(.{20,200})', False),
    ]
    
    seen_requirements = set()
    
    for pattern, is_mandatory in requirement_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            req_text = match.group(1).strip()
            # Clean up the requirement text
            req_text = re.sub(r'\s+', ' ', req_text)
            req_text = req_text.rstrip('.,;:')
            
            # Avoid duplicates
            if req_text.lower() in seen_requirements:
                continue
            seen_requirements.add(req_text.lower())
            
            # Categorize requirement
            category = categorize_requirement(req_text)
            
            requirements.append(TenderRequirement(
                requirement_id=f"REQ-{req_id:03d}",
                description=req_text,
                category=category,
                is_mandatory=is_mandatory,
                source_text=match.group(0)
            ))
            req_id += 1
    
    return requirements


def categorize_requirement(text: str) -> str:
    """Categorize a requirement based on its content."""
    text_lower = text.lower()
    
    categories = {
        "technical": ["technical", "specification", "system", "software", "hardware", "implementation"],
        "financial": ["price", "cost", "payment", "invoice", "budget", "financial"],
        "legal": ["compliance", "regulation", "law", "legal", "contract", "liability"],
        "qualification": ["experience", "qualification", "certification", "license", "capacity"],
        "timeline": ["deadline", "schedule", "timeline", "duration", "delivery", "milestone"],
        "documentation": ["document", "report", "certificate", "proposal", "submission"],
        "personnel": ["staff", "team", "personnel", "resource", "employee", "manager"],
    }
    
    for category, keywords in categories.items():
        if any(kw in text_lower for kw in keywords):
            return category
    
    return "general"


def extract_deadlines(text: str) -> list[dict]:
    """Extract deadlines and important dates from text."""
    deadlines = []
    
    # Date patterns
    date_patterns = [
        r'(\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4})',
        r'(\d{1,2}\s+(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4})',
        r'((?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4})',
    ]
    
    deadline_keywords = [
        "deadline", "due date", "submission date", "closing date", 
        "by date", "no later than", "before", "must be received"
    ]
    
    for pattern in date_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            date_str = match.group(1)
            # Find context around the date
            start = max(0, match.start() - 100)
            end = min(len(text), match.end() + 50)
            context = text[start:end]
            
            # Check if this is a deadline
            is_deadline = any(kw in context.lower() for kw in deadline_keywords)
            
            if is_deadline:
                deadlines.append({
                    "date": date_str,
                    "context": context.strip(),
                    "type": "deadline"
                })
    
    return deadlines


def extract_contacts(text: str) -> list[dict]:
    """Extract contact information from text."""
    contacts = []
    
    # Email pattern
    email_pattern = r'[\w\.-]+@[\w\.-]+\.\w+'
    emails = re.findall(email_pattern, text)
    
    for email in set(emails):
        contacts.append({
            "type": "email",
            "value": email
        })
    
    # Phone pattern
    phone_pattern = r'(?:\+\d{1,3}[-.\s]?)?\(?\d{2,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4}'
    phones = re.findall(phone_pattern, text)
    
    for phone in set(phones):
        if len(phone.replace(" ", "").replace("-", "").replace(".", "")) >= 7:
            contacts.append({
                "type": "phone",
                "value": phone.strip()
            })
    
    return contacts


def extract_evaluation_criteria(text: str) -> list[dict]:
    """Extract evaluation criteria from tender document."""
    criteria = []
    
    # Look for evaluation/scoring sections
    eval_patterns = [
        r'(?:evaluation|scoring|assessment)\s+criteria[:\s]*(.{50,500})',
        r'(?:will be evaluated|will be scored|will be assessed)\s+(?:based on|according to)[:\s]*(.{50,300})',
        r'(?:points?|score|weight(?:ing)?)[:\s]*(\d+(?:\s*%|\s*percent)?)',
    ]
    
    for pattern in eval_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            criteria.append({
                "description": match.group(1).strip() if match.group(1) else match.group(0),
                "source": match.group(0)
            })
    
    return criteria


def extract_compliance_items(text: str) -> list[dict]:
    """Extract compliance/checklist items from text."""
    items = []
    
    # Checkbox patterns
    checkbox_patterns = [
        r'[\[\(][\s\-xX]?[\]\)]\s*(.{10,200})',
        r'(?:☐|☑|✓|✔|□|■)\s*(.{10,200})',
        r'(?:yes|no|n/a)[:\s]*(.{10,100})',
    ]
    
    for pattern in checkbox_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            items.append({
                "item": match.group(1).strip(),
                "completed": False
            })
    
    return items


def generate_summary(text: str, max_length: int = 500) -> str:
    """Generate a basic summary of the document."""
    # Get first meaningful paragraphs
    paragraphs = [p.strip() for p in text.split('\n\n') if p.strip() and len(p.strip()) > 50]
    
    summary_parts = []
    current_length = 0
    
    for para in paragraphs[:5]:  # Take first 5 paragraphs max
        if current_length + len(para) > max_length:
            break
        summary_parts.append(para)
        current_length += len(para)
    
    return ' '.join(summary_parts)


# ============================================================================
# Main Analysis Function
# ============================================================================

def analyze_tender_document(file_path: Path) -> AnalysisResult:
    """Perform comprehensive analysis of a tender document."""
    doc_type = detect_document_type(file_path)
    
    # Extract text based on document type
    full_text = ""
    page_count = None
    
    if doc_type == DocumentType.PDF:
        full_text, page_count = extract_text_from_pdf(file_path)
    elif doc_type == DocumentType.DOCX:
        full_text = extract_text_from_docx(file_path)
    elif doc_type == DocumentType.HTML:
        full_text = extract_text_from_html(file_path)
    elif doc_type == DocumentType.TXT:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            full_text = f.read()
    elif doc_type == DocumentType.IMAGE:
        full_text = perform_ocr(file_path)
    
    # Extract metadata
    metadata = extract_metadata(file_path)
    if page_count:
        metadata.page_count = page_count
    metadata.word_count = len(full_text.split())
    
    # Generate document ID
    doc_id = f"DOC-{metadata.checksum[:8].upper()}"
    
    # Extract tables (PDF only)
    tables = []
    if doc_type == DocumentType.PDF:
        tables = extract_tables_from_pdf(file_path)
    
    # Create analysis result
    result = AnalysisResult(
        document_id=doc_id,
        metadata=metadata,
        sections=extract_sections(full_text),
        requirements=extract_requirements(full_text),
        tables=tables,
        deadlines=extract_deadlines(full_text),
        key_contacts=extract_contacts(full_text),
        evaluation_criteria=extract_evaluation_criteria(full_text),
        compliance_items=extract_compliance_items(full_text),
        full_text=full_text,
        summary=generate_summary(full_text)
    )
    
    return result


# ============================================================================
# MCP Server Implementation
# ============================================================================

# Create the MCP server
server = Server("tender-docs-handler")


@server.list_tools()
async def list_tools() -> ListToolsResult:
    """List all available tools."""
    return ListToolsResult(tools=[
        Tool(
            name="parse_document",
            description="Parse and extract text content from a tender document. Supports PDF, DOCX, TXT, HTML, and image files. Returns the full extracted text content.",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "Path to the document file to parse"
                    },
                    "use_ocr": {
                        "type": "boolean",
                        "description": "Whether to use OCR for scanned documents (default: auto-detect)",
                        "default": False
                    }
                },
                "required": ["file_path"]
            }
        ),
        Tool(
            name="analyze_tender",
            description="Perform comprehensive analysis of a tender document. Extracts sections, requirements, deadlines, contacts, evaluation criteria, and compliance items.",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "Path to the tender document to analyze"
                    }
                },
                "required": ["file_path"]
            }
        ),
        Tool(
            name="extract_metadata",
            description="Extract metadata from a document including file info, page count, author, title, and creation date.",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "Path to the document file"
                    }
                },
                "required": ["file_path"]
            }
        ),
        Tool(
            name="extract_tables",
            description="Extract tables from a PDF document. Returns structured table data with headers and rows.",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "Path to the PDF document"
                    }
                },
                "required": ["file_path"]
            }
        ),
        Tool(
            name="extract_requirements",
            description="Extract requirements and mandatory items from a tender document text.",
            inputSchema={
                "type": "object",
                "properties": {
                    "text": {
                        "type": "string",
                        "description": "The document text to extract requirements from"
                    }
                },
                "required": ["text"]
            }
        ),
        Tool(
            name="extract_sections",
            description="Extract and parse document sections based on headers and structure.",
            inputSchema={
                "type": "object",
                "properties": {
                    "text": {
                        "type": "string",
                        "description": "The document text to extract sections from"
                    }
                },
                "required": ["text"]
            }
        ),
        Tool(
            name="extract_deadlines",
            description="Extract deadlines and important dates from tender document text.",
            inputSchema={
                "type": "object",
                "properties": {
                    "text": {
                        "type": "string",
                        "description": "The document text to extract deadlines from"
                    }
                },
                "required": ["text"]
            }
        ),
        Tool(
            name="perform_ocr",
            description="Perform OCR (Optical Character Recognition) on scanned documents or images to extract text.",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "Path to the image or scanned PDF file"
                    }
                },
                "required": ["file_path"]
            }
        ),
        Tool(
            name="convert_document",
            description="Convert a document to a different format (e.g., DOCX to PDF, PDF to text).",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "Path to the source document"
                    },
                    "output_format": {
                        "type": "string",
                        "description": "Target format: 'pdf', 'txt', 'html', 'markdown'",
                        "enum": ["pdf", "txt", "html", "markdown"]
                    },
                    "output_path": {
                        "type": "string",
                        "description": "Path for the output file (optional)"
                    }
                },
                "required": ["file_path", "output_format"]
            }
        ),
        Tool(
            name="validate_document",
            description="Validate that a document meets basic requirements for tender submission (readable, not corrupted, contains text).",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "Path to the document to validate"
                    }
                },
                "required": ["file_path"]
            }
        ),
        Tool(
            name="compare_documents",
            description="Compare two documents and identify differences in content, structure, or requirements.",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path_1": {
                        "type": "string",
                        "description": "Path to the first document"
                    },
                    "file_path_2": {
                        "type": "string",
                        "description": "Path to the second document"
                    }
                },
                "required": ["file_path_1", "file_path_2"]
            }
        ),
        Tool(
            name="search_document",
            description="Search for specific text or patterns within a document.",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "Path to the document to search"
                    },
                    "query": {
                        "type": "string",
                        "description": "Search query or pattern"
                    },
                    "case_sensitive": {
                        "type": "boolean",
                        "description": "Whether search should be case-sensitive",
                        "default": False
                    }
                },
                "required": ["file_path", "query"]
            }
        ),
        Tool(
            name="get_document_structure",
            description="Get the hierarchical structure of a document (table of contents, section hierarchy).",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "Path to the document"
                    }
                },
                "required": ["file_path"]
            }
        ),
        Tool(
            name="list_uploaded_documents",
            description="List all documents currently uploaded and available for processing.",
            inputSchema={
                "type": "object",
                "properties": {
                    "directory": {
                        "type": "string",
                        "description": "Directory to list documents from (optional, defaults to upload directory)"
                    }
                }
            }
        ),
        Tool(
            name="import_document",
            description="Import a document from a URL or base64-encoded content into the processing directory.",
            inputSchema={
                "type": "object",
                "properties": {
                    "source": {
                        "type": "string",
                        "description": "URL or base64-encoded document content"
                    },
                    "filename": {
                        "type": "string",
                        "description": "Name to save the file as"
                    },
                    "source_type": {
                        "type": "string",
                        "description": "Type of source: 'url' or 'base64'",
                        "enum": ["url", "base64"]
                    }
                },
                "required": ["source", "filename", "source_type"]
            }
        ),
    ])


@server.call_tool()
async def call_tool(name: str, arguments: dict) -> CallToolResult:
    """Handle tool calls."""
    try:
        if name == "parse_document":
            file_path = Path(arguments["file_path"])
            use_ocr = arguments.get("use_ocr", False)
            
            if not file_path.exists():
                return CallToolResult(content=[TextContent(
                    type="text",
                    text=json.dumps({"error": f"File not found: {file_path}"})
                )])
            
            doc_type = detect_document_type(file_path)
            
            if use_ocr or doc_type == DocumentType.IMAGE:
                text = perform_ocr(file_path)
            elif doc_type == DocumentType.PDF:
                text, _ = extract_text_from_pdf(file_path)
            elif doc_type == DocumentType.DOCX:
                text = extract_text_from_docx(file_path)
            elif doc_type == DocumentType.HTML:
                text = extract_text_from_html(file_path)
            elif doc_type == DocumentType.TXT:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
            else:
                text = ""
            
            return CallToolResult(content=[TextContent(
                type="text",
                text=json.dumps({
                    "file_path": str(file_path),
                    "document_type": doc_type.value,
                    "text_length": len(text),
                    "content": text
                }, indent=2)
            )])
        
        elif name == "analyze_tender":
            file_path = Path(arguments["file_path"])
            
            if not file_path.exists():
                return CallToolResult(content=[TextContent(
                    type="text",
                    text=json.dumps({"error": f"File not found: {file_path}"})
                )])
            
            result = analyze_tender_document(file_path)
            
            return CallToolResult(content=[TextContent(
                type="text",
                text=json.dumps(result.to_dict(), indent=2)
            )])
        
        elif name == "extract_metadata":
            file_path = Path(arguments["file_path"])
            
            if not file_path.exists():
                return CallToolResult(content=[TextContent(
                    type="text",
                    text=json.dumps({"error": f"File not found: {file_path}"})
                )])
            
            metadata = extract_metadata(file_path)
            
            return CallToolResult(content=[TextContent(
                type="text",
                text=json.dumps(metadata.to_dict(), indent=2)
            )])
        
        elif name == "extract_tables":
            file_path = Path(arguments["file_path"])
            
            if not file_path.exists():
                return CallToolResult(content=[TextContent(
                    type="text",
                    text=json.dumps({"error": f"File not found: {file_path}"})
                )])
            
            tables = extract_tables_from_pdf(file_path)
            
            return CallToolResult(content=[TextContent(
                type="text",
                text=json.dumps({"tables": tables, "count": len(tables)}, indent=2)
            )])
        
        elif name == "extract_requirements":
            text = arguments["text"]
            requirements = extract_requirements(text)
            
            return CallToolResult(content=[TextContent(
                type="text",
                text=json.dumps({
                    "requirements": [r.to_dict() for r in requirements],
                    "count": len(requirements)
                }, indent=2)
            )])
        
        elif name == "extract_sections":
            text = arguments["text"]
            sections = extract_sections(text)
            
            return CallToolResult(content=[TextContent(
                type="text",
                text=json.dumps({
                    "sections": [s.to_dict() for s in sections],
                    "count": len(sections)
                }, indent=2)
            )])
        
        elif name == "extract_deadlines":
            text = arguments["text"]
            deadlines = extract_deadlines(text)
            
            return CallToolResult(content=[TextContent(
                type="text",
                text=json.dumps({
                    "deadlines": deadlines,
                    "count": len(deadlines)
                }, indent=2)
            )])
        
        elif name == "perform_ocr":
            file_path = Path(arguments["file_path"])
            
            if not file_path.exists():
                return CallToolResult(content=[TextContent(
                    type="text",
                    text=json.dumps({"error": f"File not found: {file_path}"})
                )])
            
            text = perform_ocr(file_path)
            
            return CallToolResult(content=[TextContent(
                type="text",
                text=json.dumps({
                    "file_path": str(file_path),
                    "ocr_text": text,
                    "text_length": len(text)
                }, indent=2)
            )])
        
        elif name == "convert_document":
            file_path = Path(arguments["file_path"])
            output_format = arguments["output_format"]
            output_path = arguments.get("output_path")
            
            if not file_path.exists():
                return CallToolResult(content=[TextContent(
                    type="text",
                    text=json.dumps({"error": f"File not found: {file_path}"})
                )])
            
            if not output_path:
                output_path = PROCESSED_DIR / f"{file_path.stem}.{output_format}"
            else:
                output_path = Path(output_path)
            
            try:
                if output_format == "pdf":
                    subprocess.run([
                        "soffice", "--headless", "--convert-to", "pdf",
                        "--outdir", str(output_path.parent), str(file_path)
                    ], check=True)
                elif output_format in ["txt", "markdown"]:
                    format_name = "plain" if output_format == "txt" else "markdown"
                    subprocess.run([
                        "pandoc", str(file_path), "-t", format_name, "-o", str(output_path)
                    ], check=True)
                elif output_format == "html":
                    subprocess.run([
                        "pandoc", str(file_path), "-t", "html", "-o", str(output_path)
                    ], check=True)
                
                return CallToolResult(content=[TextContent(
                    type="text",
                    text=json.dumps({
                        "success": True,
                        "input_path": str(file_path),
                        "output_path": str(output_path),
                        "format": output_format
                    }, indent=2)
                )])
            except Exception as e:
                return CallToolResult(content=[TextContent(
                    type="text",
                    text=json.dumps({"error": str(e)})
                )])
        
        elif name == "validate_document":
            file_path = Path(arguments["file_path"])
            
            validation_result = {
                "file_path": str(file_path),
                "exists": file_path.exists(),
                "readable": False,
                "has_content": False,
                "issues": []
            }
            
            if not file_path.exists():
                validation_result["issues"].append("File does not exist")
            else:
                try:
                    doc_type = detect_document_type(file_path)
                    validation_result["readable"] = True
                    validation_result["document_type"] = doc_type.value
                    
                    # Try to extract some text
                    if doc_type == DocumentType.PDF:
                        text, _ = extract_text_from_pdf(file_path)
                    elif doc_type == DocumentType.DOCX:
                        text = extract_text_from_docx(file_path)
                    elif doc_type == DocumentType.TXT:
                        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                            text = f.read()
                    else:
                        text = ""
                    
                    validation_result["has_content"] = len(text.strip()) > 0
                    validation_result["word_count"] = len(text.split())
                    
                    if not validation_result["has_content"]:
                        validation_result["issues"].append("Document appears to be empty or image-only (may need OCR)")
                    
                except Exception as e:
                    validation_result["issues"].append(f"Error reading document: {str(e)}")
            
            validation_result["valid"] = len(validation_result["issues"]) == 0
            
            return CallToolResult(content=[TextContent(
                type="text",
                text=json.dumps(validation_result, indent=2)
            )])
        
        elif name == "compare_documents":
            file_path_1 = Path(arguments["file_path_1"])
            file_path_2 = Path(arguments["file_path_2"])
            
            # Extract text from both documents
            if not file_path_1.exists() or not file_path_2.exists():
                return CallToolResult(content=[TextContent(
                    type="text",
                    text=json.dumps({"error": "One or both files not found"})
                )])
            
            result1 = analyze_tender_document(file_path_1)
            result2 = analyze_tender_document(file_path_2)
            
            comparison = {
                "document_1": {
                    "path": str(file_path_1),
                    "word_count": result1.metadata.word_count,
                    "sections_count": len(result1.sections),
                    "requirements_count": len(result1.requirements)
                },
                "document_2": {
                    "path": str(file_path_2),
                    "word_count": result2.metadata.word_count,
                    "sections_count": len(result2.sections),
                    "requirements_count": len(result2.requirements)
                },
                "differences": {
                    "word_count_diff": abs(result1.metadata.word_count - result2.metadata.word_count),
                    "sections_diff": abs(len(result1.sections) - len(result2.sections)),
                    "requirements_diff": abs(len(result1.requirements) - len(result2.requirements))
                }
            }
            
            return CallToolResult(content=[TextContent(
                type="text",
                text=json.dumps(comparison, indent=2)
            )])
        
        elif name == "search_document":
            file_path = Path(arguments["file_path"])
            query = arguments["query"]
            case_sensitive = arguments.get("case_sensitive", False)
            
            if not file_path.exists():
                return CallToolResult(content=[TextContent(
                    type="text",
                    text=json.dumps({"error": f"File not found: {file_path}"})
                )])
            
            # Extract text
            doc_type = detect_document_type(file_path)
            if doc_type == DocumentType.PDF:
                text, _ = extract_text_from_pdf(file_path)
            elif doc_type == DocumentType.DOCX:
                text = extract_text_from_docx(file_path)
            elif doc_type == DocumentType.TXT:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
            else:
                text = ""
            
            # Search
            if not case_sensitive:
                search_text = text.lower()
                search_query = query.lower()
            else:
                search_text = text
                search_query = query
            
            matches = []
            start = 0
            while True:
                pos = search_text.find(search_query, start)
                if pos == -1:
                    break
                
                # Get context around match
                context_start = max(0, pos - 50)
                context_end = min(len(text), pos + len(query) + 50)
                
                matches.append({
                    "position": pos,
                    "context": text[context_start:context_end]
                })
                start = pos + 1
            
            return CallToolResult(content=[TextContent(
                type="text",
                text=json.dumps({
                    "query": query,
                    "matches_found": len(matches),
                    "matches": matches[:50]  # Limit to first 50 matches
                }, indent=2)
            )])
        
        elif name == "get_document_structure":
            file_path = Path(arguments["file_path"])
            
            if not file_path.exists():
                return CallToolResult(content=[TextContent(
                    type="text",
                    text=json.dumps({"error": f"File not found: {file_path}"})
                )])
            
            result = analyze_tender_document(file_path)
            
            structure = {
                "document": str(file_path),
                "sections": [
                    {
                        "number": s.section_number,
                        "title": s.title,
                        "level": s.level,
                        "content_preview": s.content[:100] + "..." if len(s.content) > 100 else s.content
                    }
                    for s in result.sections
                ],
                "table_of_contents": [
                    {"level": s.level, "title": f"{s.section_number} {s.title}".strip()}
                    for s in result.sections
                ]
            }
            
            return CallToolResult(content=[TextContent(
                type="text",
                text=json.dumps(structure, indent=2)
            )])
        
        elif name == "list_uploaded_documents":
            directory = arguments.get("directory", str(UPLOAD_DIR))
            dir_path = Path(directory)
            
            if not dir_path.exists():
                return CallToolResult(content=[TextContent(
                    type="text",
                    text=json.dumps({"error": f"Directory not found: {directory}"})
                )])
            
            documents = []
            for file_path in dir_path.iterdir():
                if file_path.is_file():
                    doc_type = detect_document_type(file_path)
                    stat = file_path.stat()
                    documents.append({
                        "name": file_path.name,
                        "path": str(file_path),
                        "type": doc_type.value,
                        "size": stat.st_size,
                        "modified": datetime.fromtimestamp(stat.st_mtime).isoformat()
                    })
            
            return CallToolResult(content=[TextContent(
                type="text",
                text=json.dumps({
                    "directory": str(dir_path),
                    "documents": documents,
                    "count": len(documents)
                }, indent=2)
            )])
        
        elif name == "import_document":
            source = arguments["source"]
            filename = arguments["filename"]
            source_type = arguments["source_type"]
            
            output_path = UPLOAD_DIR / filename
            
            try:
                if source_type == "url":
                    import urllib.request
                    urllib.request.urlretrieve(source, output_path)
                elif source_type == "base64":
                    content = base64.b64decode(source)
                    with open(output_path, "wb") as f:
                        f.write(content)
                
                return CallToolResult(content=[TextContent(
                    type="text",
                    text=json.dumps({
                        "success": True,
                        "file_path": str(output_path),
                        "file_size": output_path.stat().st_size
                    }, indent=2)
                )])
            except Exception as e:
                return CallToolResult(content=[TextContent(
                    type="text",
                    text=json.dumps({"error": str(e)})
                )])
        
        else:
            return CallToolResult(content=[TextContent(
                type="text",
                text=json.dumps({"error": f"Unknown tool: {name}"})
            )])
    
    except Exception as e:
        logger.exception(f"Error in tool {name}")
        return CallToolResult(content=[TextContent(
            type="text",
            text=json.dumps({"error": str(e)})
        )])


# ============================================================================
# Main Entry Point
# ============================================================================

async def main():
    """Run the MCP server."""
    logger.info("Starting Tender Documents Handler MCP Server...")
    async with stdio_server() as (read_stream, write_stream):
        await server.run(read_stream, write_stream, server.create_initialization_options())


if __name__ == "__main__":
    asyncio.run(main())
