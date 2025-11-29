"""
Tender Documents Handler MCP Server
====================================
A FastMCP-based Model Context Protocol server for parsing, analyzing,
importing, OCR processing, and managing tender documents for tri-tender.

Compatible with FastMCP Cloud deployment.

Usage:
    Local: fastmcp run server.py
    Cloud: Deploy to fastmcp.cloud with entrypoint "server.py:mcp"
"""

import base64
import hashlib
import io
import json
import os
import re
import subprocess
import tempfile
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Optional

from fastmcp import FastMCP, Context

# ============================================================================
# FastMCP Server Instance
# ============================================================================

mcp = FastMCP(
    name="Tender Documents Handler",
    instructions="""
    This MCP server provides comprehensive document handling capabilities for tender documents.
    
    Key capabilities:
    - Parse and extract text from PDF, DOCX, TXT, HTML files
    - Perform OCR on scanned documents and images
    - Analyze tender documents to extract requirements, deadlines, and evaluation criteria
    - Extract tables from PDF documents
    - Validate document integrity and readability
    - Search within documents
    - Convert between document formats
    
    Use the 'analyze_tender' tool for comprehensive tender document analysis.
    Use 'parse_document' for simple text extraction.
    Use 'perform_ocr' for scanned documents.
    """,
    dependencies=[
        "pypdf",
        "pdfplumber", 
        "python-docx",
        "beautifulsoup4",
        "lxml",
        "pytesseract",
        "pdf2image",
        "Pillow",
    ]
)

# ============================================================================
# Configuration
# ============================================================================

# Directory configuration - uses temp directories for cloud compatibility
UPLOAD_DIR = Path(os.environ.get("TENDER_UPLOAD_DIR", tempfile.gettempdir() + "/tender-uploads"))
PROCESSED_DIR = Path(os.environ.get("TENDER_PROCESSED_DIR", tempfile.gettempdir() + "/tender-processed"))
CACHE_DIR = Path(os.environ.get("TENDER_CACHE_DIR", tempfile.gettempdir() + "/tender-cache"))

# Ensure directories exist
for dir_path in [UPLOAD_DIR, PROCESSED_DIR, CACHE_DIR]:
    dir_path.mkdir(parents=True, exist_ok=True)


# ============================================================================
# Enums and Data Classes
# ============================================================================

class DocumentType(Enum):
    PDF = "pdf"
    DOCX = "docx"
    DOC = "doc"
    TXT = "txt"
    HTML = "html"
    RTF = "rtf"
    ODT = "odt"
    IMAGE = "image"
    UNKNOWN = "unknown"


@dataclass
class DocumentMetadata:
    """Metadata extracted from a document."""
    file_name: str
    file_path: str
    file_size: int
    file_type: str
    page_count: Optional[int] = None
    title: Optional[str] = None
    author: Optional[str] = None
    created_date: Optional[str] = None
    modified_date: Optional[str] = None
    word_count: Optional[int] = None
    checksum: Optional[str] = None

    def to_dict(self) -> dict:
        return {k: v for k, v in self.__dict__.items() if v is not None}


@dataclass
class ExtractedSection:
    """A section extracted from a document."""
    title: str
    content: str
    page_number: Optional[int] = None
    section_number: Optional[str] = None
    level: int = 1

    def to_dict(self) -> dict:
        return self.__dict__


@dataclass
class TenderRequirement:
    """A requirement extracted from a tender document."""
    requirement_id: str
    description: str
    category: str
    is_mandatory: bool = False
    page_number: Optional[int] = None
    source_text: Optional[str] = None

    def to_dict(self) -> dict:
        return self.__dict__


@dataclass
class AnalysisResult:
    """Result of comprehensive tender document analysis."""
    document_id: str
    metadata: DocumentMetadata
    sections: list = field(default_factory=list)
    requirements: list = field(default_factory=list)
    tables: list = field(default_factory=list)
    deadlines: list = field(default_factory=list)
    key_contacts: list = field(default_factory=list)
    evaluation_criteria: list = field(default_factory=list)
    compliance_items: list = field(default_factory=list)
    full_text: str = ""
    summary: str = ""

    def to_dict(self) -> dict:
        return {
            "document_id": self.document_id,
            "metadata": self.metadata.to_dict(),
            "sections": [s.to_dict() if hasattr(s, 'to_dict') else s for s in self.sections],
            "requirements": [r.to_dict() if hasattr(r, 'to_dict') else r for r in self.requirements],
            "tables": self.tables,
            "deadlines": self.deadlines,
            "key_contacts": self.key_contacts,
            "evaluation_criteria": self.evaluation_criteria,
            "compliance_items": self.compliance_items,
            "full_text": self.full_text[:5000] + "..." if len(self.full_text) > 5000 else self.full_text,
            "summary": self.summary,
        }


# ============================================================================
# Core Document Processing Functions
# ============================================================================

def detect_document_type(file_path: Path) -> DocumentType:
    """Detect the type of a document based on extension."""
    extension = file_path.suffix.lower()
    
    type_mapping = {
        ".pdf": DocumentType.PDF,
        ".docx": DocumentType.DOCX,
        ".doc": DocumentType.DOC,
        ".txt": DocumentType.TXT,
        ".html": DocumentType.HTML,
        ".htm": DocumentType.HTML,
        ".rtf": DocumentType.RTF,
        ".odt": DocumentType.ODT,
        ".png": DocumentType.IMAGE,
        ".jpg": DocumentType.IMAGE,
        ".jpeg": DocumentType.IMAGE,
        ".tiff": DocumentType.IMAGE,
        ".tif": DocumentType.IMAGE,
        ".bmp": DocumentType.IMAGE,
    }
    
    return type_mapping.get(extension, DocumentType.UNKNOWN)


def calculate_checksum(file_path: Path) -> str:
    """Calculate MD5 checksum of a file."""
    hash_md5 = hashlib.md5()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_md5.update(chunk)
    return hash_md5.hexdigest()


def extract_text_from_pdf(file_path: Path) -> tuple[str, int]:
    """Extract text from a PDF file."""
    try:
        import pdfplumber
        
        text_parts = []
        page_count = 0
        
        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        return "\n\n".join(text_parts), page_count
    except ImportError:
        from pypdf import PdfReader
        
        reader = PdfReader(file_path)
        text_parts = []
        for page in reader.pages:
            text_parts.append(page.extract_text())
        
        return "\n\n".join(text_parts), len(reader.pages)


def extract_text_from_docx(file_path: Path) -> str:
    """Extract text from a DOCX file."""
    try:
        result = subprocess.run(
            ["pandoc", str(file_path), "-t", "plain"],
            capture_output=True,
            text=True,
            check=True,
            timeout=60
        )
        return result.stdout
    except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired):
        try:
            from docx import Document
            doc = Document(file_path)
            return "\n\n".join([para.text for para in doc.paragraphs])
        except ImportError:
            return ""


def extract_text_from_html(file_path: Path) -> str:
    """Extract text from an HTML file."""
    try:
        from bs4 import BeautifulSoup
        
        with open(file_path, "r", encoding="utf-8") as f:
            soup = BeautifulSoup(f.read(), "html.parser")
            for script in soup(["script", "style"]):
                script.decompose()
            return soup.get_text(separator="\n")
    except ImportError:
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
            content = re.sub(r'<[^>]+>', '', content)
            return content


def perform_ocr_on_file(file_path: Path) -> str:
    """Perform OCR on an image or scanned PDF."""
    try:
        import pytesseract
        from PIL import Image
        
        doc_type = detect_document_type(file_path)
        
        if doc_type == DocumentType.IMAGE:
            image = Image.open(file_path)
            return pytesseract.image_to_string(image)
        elif doc_type == DocumentType.PDF:
            from pdf2image import convert_from_path
            
            images = convert_from_path(file_path)
            text_parts = []
            for i, image in enumerate(images):
                text_parts.append(f"--- Page {i+1} ---\n")
                text_parts.append(pytesseract.image_to_string(image))
            return "\n".join(text_parts)
        else:
            return ""
    except ImportError as e:
        return f"OCR dependencies not available: {e}"


def extract_tables_from_pdf(file_path: Path) -> list[dict]:
    """Extract tables from a PDF document."""
    tables = []
    try:
        import pdfplumber
        
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                page_tables = page.extract_tables()
                for table_idx, table in enumerate(page_tables):
                    if table and len(table) > 1:
                        headers = table[0] if table[0] else [f"Col_{i}" for i in range(len(table[1]))]
                        rows = table[1:]
                        tables.append({
                            "page": page_num,
                            "table_index": table_idx,
                            "headers": headers,
                            "rows": rows,
                            "row_count": len(rows)
                        })
    except Exception as e:
        tables.append({"error": str(e)})
    
    return tables


def extract_metadata_from_file(file_path: Path) -> DocumentMetadata:
    """Extract metadata from a document."""
    file_stat = file_path.stat()
    doc_type = detect_document_type(file_path)
    
    metadata = DocumentMetadata(
        file_name=file_path.name,
        file_path=str(file_path),
        file_size=file_stat.st_size,
        file_type=doc_type.value,
        checksum=calculate_checksum(file_path),
        modified_date=datetime.fromtimestamp(file_stat.st_mtime).isoformat()
    )
    
    if doc_type == DocumentType.PDF:
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            metadata.page_count = len(reader.pages)
            if reader.metadata:
                metadata.title = reader.metadata.title
                metadata.author = reader.metadata.author
                if reader.metadata.creation_date:
                    metadata.created_date = str(reader.metadata.creation_date)
        except Exception:
            pass
    
    elif doc_type == DocumentType.DOCX:
        try:
            from docx import Document
            doc = Document(file_path)
            core_props = doc.core_properties
            metadata.title = core_props.title
            metadata.author = core_props.author
            if core_props.created:
                metadata.created_date = core_props.created.isoformat()
        except Exception:
            pass
    
    return metadata


# ============================================================================
# Analysis Functions
# ============================================================================

def extract_sections(text: str) -> list[ExtractedSection]:
    """Extract sections from document text."""
    sections = []
    
    patterns = [
        r'^(\d+(?:\.\d+)*\.?\s*)([A-Z][^.\n]{2,}[^\n]*)',
        r'^([A-Z]\.\s*)([A-Z][^.\n]{2,}[^\n]*)',
        r'^((?:ARTICLE|SECTION|PART)\s+[\dIVXivx]+[:\.\s]*)(.*?)$',
        r'^([A-Z][A-Z\s]{4,})$',
    ]
    
    lines = text.split('\n')
    current_section = None
    current_content = []
    
    for line in lines:
        is_header = False
        for pattern in patterns:
            match = re.match(pattern, line.strip(), re.MULTILINE)
            if match:
                if current_section:
                    current_section.content = '\n'.join(current_content).strip()
                    if current_section.content:
                        sections.append(current_section)
                
                section_num = match.group(1).strip() if len(match.groups()) > 1 else ""
                title = match.group(2).strip() if len(match.groups()) > 1 else match.group(1).strip()
                
                current_section = ExtractedSection(
                    title=title,
                    content="",
                    section_number=section_num,
                    level=1 if '.' not in section_num else section_num.count('.') + 1
                )
                current_content = []
                is_header = True
                break
        
        if not is_header and current_section:
            current_content.append(line)
    
    if current_section:
        current_section.content = '\n'.join(current_content).strip()
        if current_section.content:
            sections.append(current_section)
    
    return sections


def extract_requirements_from_text(text: str) -> list[TenderRequirement]:
    """Extract requirements from tender document text."""
    requirements = []
    req_id = 1
    
    requirement_patterns = [
        (r'(?:must|shall|should|required to|is required|mandatory)\s+(.{20,200})', True),
        (r'(?:the contractor|the bidder|the vendor|the supplier)\s+(?:must|shall|should)\s+(.{20,200})', True),
        (r'(?:minimum requirement|essential requirement|mandatory requirement)[:\s]+(.{20,200})', True),
        (r'(?:criterion|criteria)[:\s]+(.{20,200})', False),
    ]
    
    seen_requirements = set()
    
    for pattern, is_mandatory in requirement_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            req_text = match.group(1).strip()
            req_text = re.sub(r'\s+', ' ', req_text)
            req_text = req_text.rstrip('.,;:')
            
            if req_text.lower() in seen_requirements:
                continue
            seen_requirements.add(req_text.lower())
            
            category = categorize_requirement(req_text)
            
            requirements.append(TenderRequirement(
                requirement_id=f"REQ-{req_id:03d}",
                description=req_text,
                category=category,
                is_mandatory=is_mandatory,
                source_text=match.group(0)
            ))
            req_id += 1
    
    return requirements


def categorize_requirement(text: str) -> str:
    """Categorize a requirement based on its content."""
    text_lower = text.lower()
    
    categories = {
        "technical": ["technical", "specification", "system", "software", "hardware", "implementation"],
        "financial": ["price", "cost", "payment", "invoice", "budget", "financial"],
        "legal": ["compliance", "regulation", "law", "legal", "contract", "liability"],
        "qualification": ["experience", "qualification", "certification", "license", "capacity"],
        "timeline": ["deadline", "schedule", "timeline", "duration", "delivery", "milestone"],
        "documentation": ["document", "report", "certificate", "proposal", "submission"],
        "personnel": ["staff", "team", "personnel", "resource", "employee", "manager"],
    }
    
    for category, keywords in categories.items():
        if any(kw in text_lower for kw in keywords):
            return category
    
    return "general"


def extract_deadlines_from_text(text: str) -> list[dict]:
    """Extract deadlines and important dates from text."""
    deadlines = []
    
    date_patterns = [
        r'(\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4})',
        r'(\d{1,2}\s+(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4})',
        r'((?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4})',
    ]
    
    deadline_keywords = [
        "deadline", "due date", "submission date", "closing date",
        "by date", "no later than", "before", "must be received"
    ]
    
    for pattern in date_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            date_str = match.group(1)
            start = max(0, match.start() - 100)
            end = min(len(text), match.end() + 50)
            context = text[start:end]
            
            is_deadline = any(kw in context.lower() for kw in deadline_keywords)
            
            if is_deadline:
                deadlines.append({
                    "date": date_str,
                    "context": context.strip(),
                    "type": "deadline"
                })
    
    return deadlines


def extract_contacts_from_text(text: str) -> list[dict]:
    """Extract contact information from text."""
    contacts = []
    
    email_pattern = r'[\w\.-]+@[\w\.-]+\.\w+'
    emails = re.findall(email_pattern, text)
    
    for email in set(emails):
        contacts.append({"type": "email", "value": email})
    
    phone_pattern = r'(?:\+\d{1,3}[-.\s]?)?\(?\d{2,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4}'
    phones = re.findall(phone_pattern, text)
    
    for phone in set(phones):
        if len(phone.replace(" ", "").replace("-", "").replace(".", "")) >= 7:
            contacts.append({"type": "phone", "value": phone.strip()})
    
    return contacts


def extract_evaluation_criteria_from_text(text: str) -> list[dict]:
    """Extract evaluation criteria from tender document."""
    criteria = []
    
    eval_patterns = [
        r'(?:evaluation|scoring|assessment)\s+criteria[:\s]*(.{50,500})',
        r'(?:will be evaluated|will be scored|will be assessed)\s+(?:based on|according to)[:\s]*(.{50,300})',
        r'(?:points?|score|weight(?:ing)?)[:\s]*(\d+(?:\s*%|\s*percent)?)',
    ]
    
    for pattern in eval_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            criteria.append({
                "description": match.group(1).strip() if match.group(1) else match.group(0),
                "source": match.group(0)
            })
    
    return criteria


def extract_compliance_items_from_text(text: str) -> list[dict]:
    """Extract compliance/checklist items from text."""
    items = []
    
    checkbox_patterns = [
        r'[\[\(][\s\-xX]?[\]\)]\s*(.{10,200})',
        r'(?:☐|☑|✓|✔|□|■)\s*(.{10,200})',
        r'(?:yes|no|n/a)[:\s]*(.{10,100})',
    ]
    
    for pattern in checkbox_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            items.append({
                "item": match.group(1).strip(),
                "completed": False
            })
    
    return items


def generate_summary(text: str, max_length: int = 500) -> str:
    """Generate a basic summary of the document."""
    paragraphs = [p.strip() for p in text.split('\n\n') if p.strip() and len(p.strip()) > 50]
    
    summary_parts = []
    current_length = 0
    
    for para in paragraphs[:5]:
        if current_length + len(para) > max_length:
            break
        summary_parts.append(para)
        current_length += len(para)
    
    return ' '.join(summary_parts)


def analyze_tender_document_full(file_path: Path) -> AnalysisResult:
    """Perform comprehensive analysis of a tender document."""
    doc_type = detect_document_type(file_path)
    
    full_text = ""
    page_count = None
    
    if doc_type == DocumentType.PDF:
        full_text, page_count = extract_text_from_pdf(file_path)
    elif doc_type == DocumentType.DOCX:
        full_text = extract_text_from_docx(file_path)
    elif doc_type == DocumentType.HTML:
        full_text = extract_text_from_html(file_path)
    elif doc_type == DocumentType.TXT:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            full_text = f.read()
    elif doc_type == DocumentType.IMAGE:
        full_text = perform_ocr_on_file(file_path)
    
    metadata = extract_metadata_from_file(file_path)
    if page_count:
        metadata.page_count = page_count
    metadata.word_count = len(full_text.split())
    
    doc_id = f"DOC-{metadata.checksum[:8].upper()}"
    
    tables = []
    if doc_type == DocumentType.PDF:
        tables = extract_tables_from_pdf(file_path)
    
    result = AnalysisResult(
        document_id=doc_id,
        metadata=metadata,
        sections=extract_sections(full_text),
        requirements=extract_requirements_from_text(full_text),
        tables=tables,
        deadlines=extract_deadlines_from_text(full_text),
        key_contacts=extract_contacts_from_text(full_text),
        evaluation_criteria=extract_evaluation_criteria_from_text(full_text),
        compliance_items=extract_compliance_items_from_text(full_text),
        full_text=full_text,
        summary=generate_summary(full_text)
    )
    
    return result


# ============================================================================
# MCP Resources - Read-only data sources
# ============================================================================

@mcp.resource("tender://config/version")
def get_version() -> str:
    """Get the current version of the Tender Documents Handler."""
    return "1.0.0"


@mcp.resource("tender://config/supported-formats")
def get_supported_formats() -> dict:
    """Get list of supported document formats."""
    return {
        "documents": ["pdf", "docx", "doc", "txt", "html", "rtf", "odt"],
        "images": ["png", "jpg", "jpeg", "tiff", "tif", "bmp"],
        "ocr_supported": True
    }


@mcp.resource("tender://config/requirement-categories")
def get_requirement_categories() -> list:
    """Get list of requirement categories used for classification."""
    return [
        {"id": "technical", "description": "Technical specifications and system requirements"},
        {"id": "financial", "description": "Pricing, payment, and budget requirements"},
        {"id": "legal", "description": "Compliance and regulatory requirements"},
        {"id": "qualification", "description": "Experience and certification requirements"},
        {"id": "timeline", "description": "Deadline and schedule requirements"},
        {"id": "documentation", "description": "Document and report requirements"},
        {"id": "personnel", "description": "Staff and team requirements"},
        {"id": "general", "description": "Other general requirements"},
    ]


@mcp.resource("tender://uploads/{filename}")
def get_uploaded_document(filename: str) -> dict:
    """Get information about an uploaded document."""
    file_path = UPLOAD_DIR / filename
    if not file_path.exists():
        return {"error": f"File not found: {filename}"}
    
    metadata = extract_metadata_from_file(file_path)
    return metadata.to_dict()


# ============================================================================
# MCP Tools - Document Processing
# ============================================================================

@mcp.tool()
def parse_document(
    file_path: str,
    use_ocr: bool = False
) -> dict:
    """
    Parse and extract text content from a tender document.
    
    Supports PDF, DOCX, TXT, HTML, and image files.
    Use OCR for scanned documents or images.
    
    Args:
        file_path: Path to the document file to parse
        use_ocr: Force OCR processing for scanned documents
    
    Returns:
        Dictionary with file info, document type, and extracted text content
    """
    path = Path(file_path)
    
    if not path.exists():
        return {"error": f"File not found: {file_path}"}
    
    doc_type = detect_document_type(path)
    
    if use_ocr or doc_type == DocumentType.IMAGE:
        text = perform_ocr_on_file(path)
    elif doc_type == DocumentType.PDF:
        text, _ = extract_text_from_pdf(path)
    elif doc_type == DocumentType.DOCX:
        text = extract_text_from_docx(path)
    elif doc_type == DocumentType.HTML:
        text = extract_text_from_html(path)
    elif doc_type == DocumentType.TXT:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
    else:
        text = ""
    
    return {
        "file_path": str(path),
        "document_type": doc_type.value,
        "text_length": len(text),
        "word_count": len(text.split()),
        "content": text
    }


@mcp.tool()
def analyze_tender(file_path: str) -> dict:
    """
    Perform comprehensive analysis of a tender document.
    
    Extracts sections, requirements, deadlines, contacts, 
    evaluation criteria, and compliance items from the document.
    
    Args:
        file_path: Path to the tender document to analyze
    
    Returns:
        Comprehensive analysis including metadata, sections, requirements,
        deadlines, contacts, evaluation criteria, and summary
    """
    path = Path(file_path)
    
    if not path.exists():
        return {"error": f"File not found: {file_path}"}
    
    result = analyze_tender_document_full(path)
    return result.to_dict()


@mcp.tool()
def extract_metadata(file_path: str) -> dict:
    """
    Extract metadata from a document.
    
    Includes file info, page count, author, title, and creation date.
    
    Args:
        file_path: Path to the document file
    
    Returns:
        Document metadata including file size, type, author, title, etc.
    """
    path = Path(file_path)
    
    if not path.exists():
        return {"error": f"File not found: {file_path}"}
    
    metadata = extract_metadata_from_file(path)
    return metadata.to_dict()


@mcp.tool()
def extract_tables(file_path: str) -> dict:
    """
    Extract tables from a PDF document.
    
    Returns structured table data with headers and rows.
    
    Args:
        file_path: Path to the PDF document
    
    Returns:
        List of tables with headers, rows, and page numbers
    """
    path = Path(file_path)
    
    if not path.exists():
        return {"error": f"File not found: {file_path}"}
    
    tables = extract_tables_from_pdf(path)
    return {"tables": tables, "count": len(tables)}


@mcp.tool()
def extract_requirements(text: str) -> dict:
    """
    Extract requirements and mandatory items from tender document text.
    
    Identifies mandatory vs optional requirements and categorizes them.
    
    Args:
        text: The document text to extract requirements from
    
    Returns:
        List of requirements with IDs, categories, and mandatory flags
    """
    requirements = extract_requirements_from_text(text)
    return {
        "requirements": [r.to_dict() for r in requirements],
        "count": len(requirements),
        "mandatory_count": sum(1 for r in requirements if r.is_mandatory)
    }


@mcp.tool()
def extract_sections(text: str) -> dict:
    """
    Extract and parse document sections based on headers and structure.
    
    Identifies section hierarchy and content organization.
    
    Args:
        text: The document text to extract sections from
    
    Returns:
        List of sections with titles, content, and hierarchy levels
    """
    sections = extract_sections(text)
    return {
        "sections": [s.to_dict() for s in sections],
        "count": len(sections)
    }


@mcp.tool()
def extract_deadlines(text: str) -> dict:
    """
    Extract deadlines and important dates from tender document text.
    
    Identifies submission deadlines and other critical dates.
    
    Args:
        text: The document text to extract deadlines from
    
    Returns:
        List of deadlines with dates, context, and types
    """
    deadlines = extract_deadlines_from_text(text)
    return {"deadlines": deadlines, "count": len(deadlines)}


@mcp.tool()
def perform_ocr(file_path: str) -> dict:
    """
    Perform OCR on scanned documents or images to extract text.
    
    Supports image files (PNG, JPG, TIFF) and scanned PDFs.
    
    Args:
        file_path: Path to the image or scanned PDF file
    
    Returns:
        Extracted text from OCR processing
    """
    path = Path(file_path)
    
    if not path.exists():
        return {"error": f"File not found: {file_path}"}
    
    text = perform_ocr_on_file(path)
    
    return {
        "file_path": str(path),
        "ocr_text": text,
        "text_length": len(text),
        "word_count": len(text.split())
    }


@mcp.tool()
def validate_document(file_path: str) -> dict:
    """
    Validate that a document meets basic requirements for tender submission.
    
    Checks if document is readable, contains text, and is not corrupted.
    
    Args:
        file_path: Path to the document to validate
    
    Returns:
        Validation result with any issues found
    """
    path = Path(file_path)
    
    validation_result = {
        "file_path": str(path),
        "exists": path.exists(),
        "readable": False,
        "has_content": False,
        "issues": []
    }
    
    if not path.exists():
        validation_result["issues"].append("File does not exist")
        validation_result["valid"] = False
        return validation_result
    
    try:
        doc_type = detect_document_type(path)
        validation_result["readable"] = True
        validation_result["document_type"] = doc_type.value
        
        if doc_type == DocumentType.PDF:
            text, _ = extract_text_from_pdf(path)
        elif doc_type == DocumentType.DOCX:
            text = extract_text_from_docx(path)
        elif doc_type == DocumentType.TXT:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        else:
            text = ""
        
        validation_result["has_content"] = len(text.strip()) > 0
        validation_result["word_count"] = len(text.split())
        
        if not validation_result["has_content"]:
            validation_result["issues"].append("Document appears to be empty or image-only (may need OCR)")
        
    except Exception as e:
        validation_result["issues"].append(f"Error reading document: {str(e)}")
    
    validation_result["valid"] = len(validation_result["issues"]) == 0
    return validation_result


@mcp.tool()
def search_document(
    file_path: str,
    query: str,
    case_sensitive: bool = False
) -> dict:
    """
    Search for specific text or patterns within a document.
    
    Returns matching locations with surrounding context.
    
    Args:
        file_path: Path to the document to search
        query: Search query or pattern
        case_sensitive: Whether search should be case-sensitive
    
    Returns:
        List of matches with positions and surrounding context
    """
    path = Path(file_path)
    
    if not path.exists():
        return {"error": f"File not found: {file_path}"}
    
    doc_type = detect_document_type(path)
    if doc_type == DocumentType.PDF:
        text, _ = extract_text_from_pdf(path)
    elif doc_type == DocumentType.DOCX:
        text = extract_text_from_docx(path)
    elif doc_type == DocumentType.TXT:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
    else:
        text = ""
    
    if not case_sensitive:
        search_text = text.lower()
        search_query = query.lower()
    else:
        search_text = text
        search_query = query
    
    matches = []
    start = 0
    while True:
        pos = search_text.find(search_query, start)
        if pos == -1:
            break
        
        context_start = max(0, pos - 50)
        context_end = min(len(text), pos + len(query) + 50)
        
        matches.append({
            "position": pos,
            "context": text[context_start:context_end]
        })
        start = pos + 1
    
    return {
        "query": query,
        "matches_found": len(matches),
        "matches": matches[:50]
    }


@mcp.tool()
def get_document_structure(file_path: str) -> dict:
    """
    Get the hierarchical structure of a document.
    
    Returns table of contents and section hierarchy.
    
    Args:
        file_path: Path to the document
    
    Returns:
        Document structure with sections and hierarchy
    """
    path = Path(file_path)
    
    if not path.exists():
        return {"error": f"File not found: {file_path}"}
    
    result = analyze_tender_document_full(path)
    
    structure = {
        "document": str(path),
        "sections": [
            {
                "number": s.section_number,
                "title": s.title,
                "level": s.level,
                "content_preview": s.content[:100] + "..." if len(s.content) > 100 else s.content
            }
            for s in result.sections
        ],
        "table_of_contents": [
            {"level": s.level, "title": f"{s.section_number} {s.title}".strip()}
            for s in result.sections
        ]
    }
    
    return structure


@mcp.tool()
def import_document(
    source: str,
    filename: str,
    source_type: str = "base64"
) -> dict:
    """
    Import a document from base64-encoded content.
    
    Args:
        source: Base64-encoded document content
        filename: Name to save the file as
        source_type: Type of source - currently only 'base64' supported
    
    Returns:
        Path to imported document and file size
    """
    output_path = UPLOAD_DIR / filename
    
    try:
        if source_type == "base64":
            content = base64.b64decode(source)
            with open(output_path, "wb") as f:
                f.write(content)
        else:
            return {"error": f"Unsupported source type: {source_type}"}
        
        return {
            "success": True,
            "file_path": str(output_path),
            "file_size": output_path.stat().st_size
        }
    except Exception as e:
        return {"error": str(e)}


@mcp.tool()
def list_documents(directory: str = None) -> dict:
    """
    List all documents available for processing.
    
    Args:
        directory: Directory to list documents from (optional)
    
    Returns:
        List of documents with metadata
    """
    dir_path = Path(directory) if directory else UPLOAD_DIR
    
    if not dir_path.exists():
        return {"error": f"Directory not found: {dir_path}"}
    
    documents = []
    for file_path in dir_path.iterdir():
        if file_path.is_file():
            doc_type = detect_document_type(file_path)
            stat = file_path.stat()
            documents.append({
                "name": file_path.name,
                "path": str(file_path),
                "type": doc_type.value,
                "size": stat.st_size,
                "modified": datetime.fromtimestamp(stat.st_mtime).isoformat()
            })
    
    return {
        "directory": str(dir_path),
        "documents": documents,
        "count": len(documents)
    }


# ============================================================================
# MCP Prompts - Reusable templates for LLM interactions
# ============================================================================

@mcp.prompt()
def analyze_tender_prompt(file_path: str) -> str:
    """
    Generate a prompt for comprehensive tender document analysis.
    
    Args:
        file_path: Path to the tender document
    
    Returns:
        A structured prompt for analyzing the tender document
    """
    return f"""Please analyze the tender document at: {file_path}

Follow these steps:
1. First, use the 'analyze_tender' tool to get a comprehensive analysis
2. Review the extracted requirements, separating mandatory from optional
3. Identify all deadlines and create a timeline
4. List the evaluation criteria and their weights if available
5. Note any compliance requirements or certifications needed
6. Summarize key contacts and submission instructions

Provide a structured report with:
- Executive Summary
- Key Requirements (categorized)
- Submission Deadlines
- Evaluation Criteria
- Compliance Checklist
- Recommended Next Steps"""


@mcp.prompt()
def compare_requirements_prompt(doc1_path: str, doc2_path: str) -> str:
    """
    Generate a prompt for comparing requirements between two tender documents.
    
    Args:
        doc1_path: Path to the first tender document
        doc2_path: Path to the second tender document
    
    Returns:
        A structured prompt for comparing requirements
    """
    return f"""Compare the requirements between these two tender documents:
- Document 1: {doc1_path}
- Document 2: {doc2_path}

Steps:
1. Use 'analyze_tender' on both documents
2. Compare the extracted requirements
3. Identify common requirements
4. Highlight unique requirements in each document
5. Note differences in evaluation criteria
6. Compare deadlines

Provide a comparison report with:
- Common Requirements
- Unique to Document 1
- Unique to Document 2
- Differences in Evaluation
- Timeline Comparison"""


@mcp.prompt()
def extract_compliance_checklist_prompt(file_path: str) -> str:
    """
    Generate a prompt for creating a compliance checklist from a tender.
    
    Args:
        file_path: Path to the tender document
    
    Returns:
        A prompt for generating a compliance checklist
    """
    return f"""Create a comprehensive compliance checklist for the tender at: {file_path}

Steps:
1. Use 'analyze_tender' to extract all requirements
2. Filter for mandatory requirements
3. Identify documentation requirements
4. Note certification requirements
5. List qualification requirements

Generate a checklist with:
- [ ] Document Requirements
- [ ] Certification Requirements
- [ ] Experience/Qualification Requirements
- [ ] Technical Compliance Items
- [ ] Financial Compliance Items
- [ ] Administrative Requirements

Include the source reference for each item."""


# ============================================================================
# Entry Point
# ============================================================================

if __name__ == "__main__":
    # Local development - run with stdio transport
    mcp.run()
